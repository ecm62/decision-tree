import streamlit as st
import graphviz
import pandas as pd
from docx import Document
from docx.shared import Inches
import re
import textwrap
import io
import os

# ==========================================
# 介面基礎配置與高級 CSS 注入
# ==========================================
st.set_page_config(page_title="英俊的小羊心智圖製作", page_icon="🐏", layout="wide")

st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stApp { background-color: #f4f6f9; }
    
    div.stButton > button:first-child {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white; font-weight: 600; font-size: 16px; border-radius: 8px;
        border: none; padding: 10px 24px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }
    div.stButton > button:first-child:hover {
        transform: translateY(-2px); box-shadow: 0 6px 12px rgba(0,0,0,0.2);
    }
    
    div[data-testid="stDownloadButton"] > button {
        border: 1px solid #c0ccda; background-color: #ffffff; color: #1e3c72;
        border-radius: 6px; transition: all 0.2s ease;
    }
    div[data-testid="stDownloadButton"] > button:hover {
        border-color: #2a5298; color: #2a5298; background-color: #f0f4f8;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 核心解析引擎庫
# ==========================================
def parse_indentation(text):
    nodes_dict, edges, stack, node_counter = {}, [], [], 0
    def get_id():
        nonlocal node_counter
        node_counter += 1
        return f"IND_{node_counter}"

    for line in text.strip().split('\n'):
        if not line.strip(): continue
        level = len(line) - len(line.lstrip(' \t'))
        label = line.strip()
        node_id = get_id()
        nodes_dict[node_id] = {"label": label, "type": "standard", "depth": level}
        while stack and stack[-1][0] >= level: stack.pop()
        if stack: edges.append((stack[-1][1], node_id))
        stack.append((level, node_id))
    return nodes_dict, edges

def parse_mermaid(text):
    nodes_dict, edges = {}, []
    node_pattern = re.compile(r'([A-Za-z0-9_]+)(?:\[|\(|\{)(.*?)(?:\]|\)|\})')
    for line in text.strip().split('\n'):
        line = line.strip()
        if not line or line.startswith('%%') or line.startswith('graph'): continue
        if '-->' in line:
            parts = [p.strip() for p in line.split('-->')]
            chain_ids = []
            for part in parts:
                match = node_pattern.search(part)
                if match:
                    node_id = match.group(1).strip()
                    label = match.group(2).replace('<br>', '\n').strip()
                    nodes_dict[node_id] = {"label": label, "type": "standard", "depth": 0}
                    chain_ids.append(node_id)
                else:
                    node_id = part.strip()
                    if node_id:
                        if node_id not in nodes_dict: nodes_dict[node_id] = {"label": node_id, "type": "standard", "depth": 0}
                        chain_ids.append(node_id)
            for i in range(len(chain_ids) - 1): edges.append((chain_ids[i], chain_ids[i+1]))
    return nodes_dict, edges

def parse_arrow_chain(text):
    nodes_dict, edges, label_to_id, node_counter = {}, [], {}, 0

    def get_or_create_node(label, depth):
        nonlocal node_counter
        label = label.strip()
        if label not in label_to_id:
            node_counter += 1
            node_id = f"ARR_{node_counter}"
            label_to_id[label] = node_id
            
            node_type = "standard"
            disease_kws = ["病", "炎", "症", "感染", "osis", "itis", "syndrome", "fever", "virus", "bacteria"]
            treatment_kws = ["治療", "首選", "次選", "1st:", "2nd:", "penicillin", "mycin", "sporin", "藥", "支持", "support", "vaccine"]
            
            label_lower = label.lower()
            if any(kw in label_lower for kw in treatment_kws): node_type = "treatment"
            elif any(kw in label_lower for kw in disease_kws): node_type = "disease"
            
            nodes_dict[node_id] = {"label": label, "type": node_type, "depth": depth}
        return label_to_id[label]

    text = text.replace('->', '➔').replace('=>', '➔').replace('➡️', '➔')
    
    for line in text.strip().split('\n'):
        line = line.strip()
        if not line: continue
        
        if '➔' in line:
            parts = line.split('➔')
            for i in range(len(parts) - 1):
                p_label, c_label = parts[i].strip(), parts[i+1].strip()
                if not p_label or not c_label: continue
                
                p_id = get_or_create_node(p_label, depth=i)
                c_id = get_or_create_node(c_label, depth=i+1)
                
                if (p_id, c_id) not in edges: edges.append((p_id, c_id))
        else: 
            get_or_create_node(line, depth=0)

    return nodes_dict, edges

def auto_detect_and_parse(text):
    if '-->' in text or 'graph LR' in text or 'graph TB' in text: return parse_mermaid(text), "Mermaid 語法模式"
    elif '➔' in text or '->' in text: return parse_arrow_chain(text), "連續推演引擎 (智能語意)"
    else: return parse_indentation(text), "空白縮排模式"

def format_label_wrap(text, width):
    lines = text.split('\n')
    return '\n'.join([textwrap.fill(line, width=width) if len(line) > width else line for line in lines])

# ==========================================
# 主程式邏輯與 UI 結構
# ==========================================
st.title("🐏 英俊的小羊 專業決策樹 Decision-tree (心智圖)生成器")
st.markdown("---")

col1, col2 = st.columns([1.2, 2.0], gap="large")

with col1:
    with st.form("main_form", border=False):
        st.markdown("#### 📥 數據輸入區")
        
        graph_title = st.text_area(
            "圖表頂部標題", 
            value="心智圖           ", 
            height=68
        )
        
        input_text = st.text_area("結構文字 (若要設定大標題，請獨立換行輸入即可)", height=250, placeholder="例如：\n呼吸道疾病類別\n豬藍耳病 ➔ 呼吸急促 ➔ 抗生素治療")
        
        with st.expander("⚙️ 進階渲染與上色參數設定", expanded=True):
            color_mode = st.selectbox("選擇上色模式", [
                "智能分類上色 (動態層級漸層版)", 
                "層級統一上色 (專業版面首選)",
                "企業冷色調 (高階 SOP 專用)",
                "高對比警戒 (異常排查與警示)",
                "極簡學術灰階 (黑白列印/論文專用)"
            ])
            
            d1, d2 = st.columns(2)
            with d1: direction = st.radio("排版方向", ["橫式 (左至右)", "直式 (上至下)"])
            with d2: line_style = st.radio("連線風格", ["直角折線", "彎曲線條"])
            
            s1, s2 = st.columns(2)
            with s1: node_shape = st.radio("節點形狀", ["方框", "圓框", "無框"])
            with s2: density = st.radio("排版密度", ["緊密", "適中", "鬆散"])
            
            font_choice = st.selectbox("字體設定", ["正黑體 (預設現代)", "明體 (正式學術)", "楷體 (傳統人文)"])
            wrap_width = st.slider("自動斷行字數限制 (字/行)", 5, 40, 15)
        
        st.write("") 
        submit_btn = st.form_submit_button("🚀 執行智能分析與生成", use_container_width=True)

with col2:
    st.markdown("#### 📊 視覺化分析結果")
    if submit_btn:
        if not input_text.strip():
            st.warning("⚠️ 請先在左側輸入文字數據。")
        else:
            with st.spinner("系統正在進行病理結構解析與專業級渲染..."):
                (nodes_dict, edges), detected_mode = auto_detect_and_parse(text=input_text)
                st.success(f"✓ 數據解析成功 | 系統判定格式：**{detected_mode}**")

                dir_map = {"橫式 (左至右)": "LR", "直式 (上至下)": "TB"}
                line_map = {"直角折線": "ortho", "彎曲線條": "spline"}
                shape_map = {"方框": "box", "圓框": "ellipse", "無框": "plaintext"}
                density_map = {
                    "緊密": {"nodesep": "0.1", "ranksep": "0.2"},
                    "適中": {"nodesep": "0.4", "ranksep": "0.5"},
                    "鬆散": {"nodesep": "0.8", "ranksep": "1.0"}
                }
                font_map = {
                    "正黑體 (預設現代)": "WenQuanYi Zen Hei",
                    "明體 (正式學術)": "AR PL UMing TW",
                    "楷體 (傳統人文)": "AR PL UKai TW"
                }

                selected_shape = shape_map[node_shape]
                selected_font = font_map[font_choice]

                # 動態產生 PS 文字說明
                if color_mode == "智能分類上色 (動態層級漸層版)":
                    legend_text = "PS. 顏色依據：[深色白字] 起始點 ｜ [紅色系] 疾病 ｜ [綠色系] 方案 ｜ [灰白漸層] 分類"
                elif color_mode == "層級統一上色 (專業版面首選)":
                    legend_text = "PS. 顏色依據：[深藍] 起始點 ｜ [綠框] 方案 ｜ [淺色] 階層節點"
                elif color_mode == "企業冷色調 (高階 SOP 專用)":
                    legend_text = "PS. 顏色依據：[深青] 起始點 ｜ [冷綠] 方案 ｜ [冷藍灰漸層] 階層節點"
                elif color_mode == "高對比警戒 (異常排查與警示)":
                    legend_text = "PS. 顏色依據：[極黑底黃字] 起始點 ｜ [警示橘] 排查節點 ｜ 適合疾病與耗損追蹤"
                elif color_mode == "極簡學術灰階 (黑白列印/論文專用)":
                    legend_text = "PS. 顏色依據：[黑底白字] 起始點 ｜ [虛線框] 方案 ｜ 確保黑白列印不失真"

                dot = graphviz.Digraph(format='png')
                
                dot.attr(
                    rankdir=dir_map[direction], 
                    splines=line_map[line_style], 
                    nodesep=density_map[density]["nodesep"], 
                    ranksep=density_map[density]["ranksep"],
                    label=legend_text,         
                    labelloc="b",              
                    fontname=selected_font,
                    fontsize="11",
                    fontcolor="#888888"
                )
                
                # HTML 標題與 LOGO 渲染引擎
                # 將純文字換行轉換為 HTML 的 <BR/>，並判斷檔案庫中是否有 logo.png
                safe_title = graph_title.replace('\n', '<BR/>')
                
                if os.path.exists("logo.png"):
                    # 若檔案存在，運用 HTML Table 技術將 LOGO 完美貼齊於標題左側
                    html_label = f'<<TABLE BORDER="0" CELLBORDER="0" CELLSPACING="0"><TR><TD><IMG SRC="logo.png"/></TD><TD ALIGN="CENTER"><FONT FACE="{selected_font}" POINT-SIZE="18" COLOR="#1e3c72"><B>{safe_title}</B></FONT></TD></TR></TABLE>>'
                else:
                    # 容錯機制：若您忘記上傳，系統自動降級為純文字標題，避免當機
                    html_label = f'<<TABLE BORDER="0" CELLBORDER="0" CELLSPACING="0"><TR><TD ALIGN="CENTER"><FONT FACE="{selected_font}" POINT-SIZE="18" COLOR="#1e3c72"><B>{safe_title}</B></FONT></TD></TR></TABLE>>'

                with dot.subgraph(name='cluster_main') as c:
                    c.attr(
                        label=html_label,
                        labelloc="t",          
                        color="none"           
                    )
                
                    for node_id, data in nodes_dict.items():
                        raw_label, node_type, depth = data["label"], data["type"], data["depth"]
                        formatted_label = format_label_wrap(raw_label, int(wrap_width))
                        
                        attrs = {"fontname": selected_font, "fontsize": "12", "color": "#555555"}
                        act_shape = selected_shape if selected_shape != "plaintext" else "box"
                        is_even_depth = (depth % 2 == 0)
                        
                        if color_mode == "智能分類上色 (動態層級漸層版)":
                            if depth == 0:
                                if node_type == "disease": attrs.update({"shape": "box", "style": "filled,rounded", "fillcolor": "#c0392b", "color": "#922b21", "fontcolor": "white", "fontweight": "bold", "fontsize": "14"})
                                elif node_type == "treatment": attrs.update({"shape": "box", "style": "filled,rounded", "fillcolor": "#27ae60", "color": "#1d8348", "fontcolor": "white", "fontweight": "bold", "fontsize": "14"})
                                else: attrs.update({"shape": "box", "style": "filled,rounded", "fillcolor": "#34495e", "color": "#2c3e50", "fontcolor": "white", "fontweight": "bold", "fontsize": "14"})
                            else:
                                if node_type == "disease":
                                    fill = "#fdedec" if is_even_depth else "#fadbd8"
                                    attrs.update({"shape": act_shape, "style": "filled,rounded", "fillcolor": fill, "color": "#e74c3c", "fontcolor": "#78281f"})
                                elif node_type == "treatment":
                                    fill = "#eafaf1" if is_even_depth else "#d5f5e3"
                                    attrs.update({"shape": act_shape, "style": "filled,rounded", "fillcolor": fill, "color": "#2ecc71", "fontcolor": "#186a3b", "fontsize": "11"})
                                else:
                                    fill = "#ebedef" if is_even_depth else "#fdfefe"
                                    attrs.update({"shape": selected_shape, "style": "filled,rounded", "fillcolor": fill, "color": "#bdc3c7", "fontcolor": "#2c3e50"})
                                    
                        elif color_mode == "層級統一上色 (專業版面首選)":
                            if depth == 0:
                                attrs.update({"shape": "box", "style": "filled,rounded", "fillcolor": "#2a5298", "fontcolor": "white", "fontsize": "14", "fontweight": "bold", "color": "#1e3c72"})
                            elif node_type == "treatment":
                                attrs.update({"shape": "box", "style": "rounded", "color": "#009900", "fontsize": "11", "fontcolor": "#006600"})
                            else:
                                bg_color = "#ffffff" if is_even_depth else "#f0f4f8"
                                attrs.update({"shape": selected_shape, "style": "filled,rounded", "fillcolor": bg_color, "color": "#a0a0a0"})
                                
                        elif color_mode == "企業冷色調 (高階 SOP 專用)":
                            if depth == 0:
                                attrs.update({"shape": "box", "style": "filled,rounded", "fillcolor": "#117864", "color": "#0b5345", "fontcolor": "white", "fontweight": "bold", "fontsize": "14"})
                            elif node_type == "treatment":
                                attrs.update({"shape": "box", "style": "rounded", "color": "#17a589", "fontsize": "11", "fontcolor": "#0e6251"})
                            else:
                                fill = "#e8f8f5" if is_even_depth else "#ebedef"
                                attrs.update({"shape": selected_shape, "style": "filled,rounded", "fillcolor": fill, "color": "#aeb6bf", "fontcolor": "#212f3d"})
                                
                        elif color_mode == "高對比警戒 (異常排查與警示)":
                            if depth == 0:
                                attrs.update({"shape": "box", "style": "filled,rounded", "fillcolor": "#17202a", "color": "#000000", "fontcolor": "#f1c40f", "fontweight": "bold", "fontsize": "14"})
                            elif node_type == "treatment":
                                attrs.update({"shape": "box", "style": "filled,rounded", "fillcolor": "#fef9e7", "color": "#f39c12", "fontcolor": "#d68910", "fontsize": "11"})
                            else:
                                fill = "#fdf2e9" if is_even_depth else "#fbeee6"
                                attrs.update({"shape": selected_shape, "style": "filled,rounded", "fillcolor": fill, "color": "#e67e22", "fontcolor": "#873600"})
                                
                        elif color_mode == "極簡學術灰階 (黑白列印/論文專用)":
                            if depth == 0:
                                attrs.update({"shape": "box", "style": "filled,rounded", "fillcolor": "#333333", "color": "#000000", "fontcolor": "white", "fontweight": "bold", "fontsize": "14"})
                            elif node_type == "treatment":
                                attrs.update({"shape": "box", "style": "dashed,rounded", "fillcolor": "none", "color": "#000000", "fontcolor": "#000000", "fontsize": "11"})
                            else:
                                fill = "#ffffff" if is_even_depth else "#f2f2f2"
                                attrs.update({"shape": selected_shape, "style": "filled,rounded", "fillcolor": fill, "color": "#666666", "fontcolor": "#000000"})
                        
                        c.node(node_id, formatted_label, **attrs)

                    for parent, child in edges: c.edge(parent, child, color="#888888")

                # 產出數據流
                png_data = dot.pipe(format='png')
                pdf_data = dot.pipe(format='pdf')
                svg_data = dot.pipe(format='svg')
                
                st.image(png_data, use_container_width=True)
                
                doc = Document()
                doc.add_heading(f'決策分析結構 ({detected_mode})', 0)
                doc.add_picture(io.BytesIO(png_data), width=Inches(6.5))
                doc.add_heading('原始數據記錄', level=1)
                for line in input_text.strip().split('\n'): doc.add_paragraph(line)
                docx_output = io.BytesIO()
                doc.save(docx_output)
                
                readable_edges = [(nodes_dict[p]["label"].replace('\n', ' '), nodes_dict[c]["label"].replace('\n', ' ')) for p, c in edges]
                df = pd.DataFrame(readable_edges, columns=["上層節點", "下層節點"])
                excel_output = io.BytesIO()
                df.to_excel(excel_output, index=False)
                
                md_data = f"# 決策結構數據 ({detected_mode})\n\n```text\n{input_text.strip()}\n```".encode('utf-8')

                st.markdown("---")
                st.markdown("#### 📤 專業報告與原始檔匯出")
                
                d_col1, d_col2, d_col3 = st.columns(3)
                d_col4, d_col5, d_col6 = st.columns(3)
                
                d_col1.download_button("📄 PDF 高清報告", data=pdf_data, file_name="Decision_Tree.pdf", mime="application/pdf", use_container_width=True)
                d_col2.download_button("🖼️ PNG 高畫質圖", data=png_data, file_name="Decision_Tree.png", mime="image/png", use_container_width=True)
                d_col3.download_button("📐 SVG 向量圖", data=svg_data, file_name="Decision_Tree.svg", mime="image/svg+xml", use_container_width=True)
                
                d_col4.download_button("📝 WORD 編輯檔", data=docx_output.getvalue(), file_name="Decision_Tree.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                d_col5.download_button("📊 EXCEL 關聯表", data=excel_output.getvalue(), file_name="Decision_Tree.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
                d_col6.download_button("💻 MD 原始文字", data=md_data, file_name="Decision_Tree.md", mime="text/markdown", use_container_width=True)
    else:
        st.info("💡 系統閒置中。請於左側輸入數據並點擊「🚀 執行智能分析與生成」按鈕以檢視圖形。")
